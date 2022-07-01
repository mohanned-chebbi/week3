import requests,nmap,re
import socket as s
from fpdf import FPDF
from docx import Document
from docx.shared import Inches

domain = input("Enter domain name: ")
ip = s.gethostbyname(domain)
printdomain=(f'the ip address of {domain} is {ip}')
print(printdomain)
nmScan = nmap.PortScanner()
nmScan.scan(ip)

for host in nmScan.all_hosts():
    for proto in nmScan[host].all_protocols():
        lport = nmScan[host][proto].keys()
        for port in lport:
            openPorts=('Open Port: '+str(port)+'/'+proto.upper()+' - Service: '+nmScan[host][proto][port]['name']+'\n')
print(openPorts)

res=nmScan.scan(hosts=ip,arguments="-sV --script=vuln ",ports=openPorts)

cvee=re.compile(r'(CVE-(\d)*-(\d)*)').findall(str(res))
"""
cveid=[]
for i in range(len(cvee)):
      cveid=(cvee[i][0])
"""
print(cvee)

file = open('subdomains_name.txt', 'r')
content = file.read()
subdomainsNames = content.splitlines()
for subdomain  in  subdomainsNames:
    url1 = f"http://{subdomain}.{domain}"
    url2 = f"https://{subdomain}.{domain}"
    try:
        requests.get(url1)
        URL1=(f'[+] {url1}')
        requests.get(url2)
        URL2=(f'[+] {url2}')
    except requests.ConnectionError:
        pass

"""
Scan= []
Scan.append(printdomain)
Scan.append(openPorts)
Scan.append(cvee)
Scan.append(URL1)
Scan.append(URL2)
#printdomain,'\n',openPorts,cvee,URL1,URL2

pdf = FPDF()
 
# Add a page
pdf.add_page()
 
# set style and size of font
# that you want in the pdf
pdf.set_font("Arial", size = 15)
 
# create a cell
pdf.cell(200, 10, txt = "Scan Report",
         ln = 1, align = 'C')
#pdf.cell(200, 10, txt = Scan )
#pdf.write(Scan)
pdf.cell(150, 10, txt = printdomain)

pdf.cell(150, 10, txt = openPorts)
#pdf.cell(150, 10, txt = cvee)
pdf.cell(150, 10, txt = URL1)
pdf.cell(150, 10, txt = URL2)
 
# save the pdf with name .pdf
pdf.output("AutomatedScanner.pdf")  
"""
document = Document()

document.add_heading('Scan Repor', 0)

p = document.add_paragraph(printdomain)
p = document.add_paragraph(openPorts)
p = document.add_paragraph(cvee)
p = document.add_paragraph(URL1)
p = document.add_paragraph(URL2)

document.add_page_break()

document.save('AutomatedScanner.docx')
